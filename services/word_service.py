"""
=========================================================
Word Service
=========================================================
Membuat dokumen Word dari hasil analisis.
=========================================================
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt

from services.models import HasilAnalisis


class WordService:
    """
    Service untuk membuat file Word hasil analisis.
    """

    def __init__(self):

        self.folder = Path("hasil")

        self.folder.mkdir(exist_ok=True)

    # =====================================================
    # EXPORT WORD
    # =====================================================

    def export(
        self,
        hasil: HasilAnalisis
    ) -> Path:

        document = Document()

        # ---------------------------------------------
        # Judul
        # ---------------------------------------------

        heading = document.add_heading(
            hasil.mata_kuliah,
            level=1
        )

        heading.style.font.size = Pt(20)

        # ---------------------------------------------
        # Informasi
        # ---------------------------------------------

        document.add_paragraph(
            f"Model AI : {hasil.model}"
        )

        document.add_paragraph(
            f"Waktu Proses : {hasil.waktu_proses:.2f} detik"
        )

        document.add_paragraph()

        # ---------------------------------------------
        # Isi Refleksi
        # ---------------------------------------------

        document.add_heading(
            "Hasil Analisis",
            level=2
        )

        document.add_paragraph(
            hasil.isi
        )

        # ---------------------------------------------
        # Simpan
        # ---------------------------------------------

        file = self.folder / f"{hasil.mata_kuliah}.docx"

        document.save(file)

        return file